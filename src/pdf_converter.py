"""
PDF to Word Converter Module
Converts PDF files to Word documents using pdf2docx library
Also supports PDF to Markdown conversion using pymupdf4llm
Also supports Word to Markdown conversion using python-docx
"""

import os
from pathlib import Path
from pdf2docx import Converter
from typing import Optional
import pymupdf4llm
from docx import Document


class PDFToWordConverter:
    """Class to handle PDF to Word conversion"""
    
    def __init__(self, input_dir: str = "input", output_dir: str = "output"):
        """
        Initialize the converter with input and output directories
        
        Args:
            input_dir: Directory containing PDF files
            output_dir: Directory to save converted Word files
        """
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        
        # Create directories if they don't exist
        self.input_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def convert_single_pdf(self, pdf_path: str, output_path: Optional[str] = None) -> bool:
        """
        Convert a single PDF file to Word document
        
        Args:
            pdf_path: Path to the PDF file
            output_path: Optional custom output path for the Word file
            
        Returns:
            bool: True if conversion successful, False otherwise
        """
        try:
            pdf_file = Path(pdf_path)
            
            if not pdf_file.exists():
                print(f"Error: PDF file not found: {pdf_path}")
                return False
            
            if not pdf_file.suffix.lower() == '.pdf':
                print(f"Error: File is not a PDF: {pdf_path}")
                return False
            
            # Generate output path if not provided
            if output_path is None:
                output_filename = pdf_file.stem + '.docx'
                output_path = self.output_dir / output_filename
            else:
                output_path = Path(output_path)
            
            print(f"Converting: {pdf_file.name}")
            print(f"Output: {output_path}")
            
            # Create converter and convert
            cv = Converter(str(pdf_file))
            cv.convert(str(output_path))
            cv.close()
            
            print(f"✓ Successfully converted: {pdf_file.name}")
            return True
            
        except Exception as e:
            print(f"✗ Error converting {pdf_path}: {str(e)}")
            return False
    
    def convert_to_markdown(self, pdf_path: str, output_path: Optional[str] = None) -> bool:
        """
        Convert a single PDF file to Markdown document
        
        Args:
            pdf_path: Path to the PDF file
            output_path: Optional custom output path for the Markdown file
            
        Returns:
            bool: True if conversion successful, False otherwise
        """
        try:
            pdf_file = Path(pdf_path)
            
            if not pdf_file.exists():
                print(f"Error: PDF file not found: {pdf_path}")
                return False
            
            if not pdf_file.suffix.lower() == '.pdf':
                print(f"Error: File is not a PDF: {pdf_path}")
                return False
            
            # Generate output path if not provided
            if output_path is None:
                output_filename = pdf_file.stem + '.md'
                output_path = self.output_dir / output_filename
            else:
                output_path = Path(output_path)
            
            print(f"Converting to Markdown: {pdf_file.name}")
            print(f"Output: {output_path}")
            
            # Convert PDF to Markdown using pymupdf4llm
            md_text = pymupdf4llm.to_markdown(str(pdf_file))
            
            # Write markdown content to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(md_text)
            
            print(f"✓ Successfully converted to Markdown: {pdf_file.name}")
            return True
            
        except Exception as e:
            print(f"✗ Error converting {pdf_path} to Markdown: {str(e)}")
            return False
    
    def convert_word_to_markdown(self, docx_path: str, output_path: Optional[str] = None) -> bool:
        """
        Convert a Word document to Markdown
        
        Args:
            docx_path: Path to the Word document (.docx)
            output_path: Optional custom output path for the Markdown file
            
        Returns:
            bool: True if conversion successful, False otherwise
        """
        try:
            docx_file = Path(docx_path)
            
            if not docx_file.exists():
                print(f"Error: Word file not found: {docx_path}")
                return False
            
            if not docx_file.suffix.lower() in ['.docx', '.doc']:
                print(f"Error: File is not a Word document: {docx_path}")
                return False
            
            # Generate output path if not provided
            if output_path is None:
                output_filename = docx_file.stem + '.md'
                output_path = self.output_dir / output_filename
            else:
                output_path = Path(output_path)
            
            print(f"Converting Word to Markdown: {docx_file.name}")
            print(f"Output: {output_path}")
            
            # Load Word document
            doc = Document(str(docx_file))
            
            # Convert to Markdown
            markdown_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    markdown_content.append("")
                    continue
                
                # Check paragraph style for headings
                style_name = paragraph.style.name.lower()
                
                if 'heading 1' in style_name:
                    markdown_content.append(f"# {text}")
                elif 'heading 2' in style_name:
                    markdown_content.append(f"## {text}")
                elif 'heading 3' in style_name:
                    markdown_content.append(f"### {text}")
                elif 'heading 4' in style_name:
                    markdown_content.append(f"#### {text}")
                elif 'heading 5' in style_name:
                    markdown_content.append(f"##### {text}")
                elif 'heading 6' in style_name:
                    markdown_content.append(f"###### {text}")
                else:
                    # Check for bold/italic formatting in runs
                    formatted_text = []
                    for run in paragraph.runs:
                        run_text = run.text
                        if run.bold and run.italic:
                            formatted_text.append(f"***{run_text}***")
                        elif run.bold:
                            formatted_text.append(f"**{run_text}**")
                        elif run.italic:
                            formatted_text.append(f"*{run_text}*")
                        else:
                            formatted_text.append(run_text)
                    
                    markdown_content.append(''.join(formatted_text))
            
            # Process tables
            for table in doc.tables:
                markdown_content.append("")  # Add blank line before table
                
                # Header row
                header_cells = [cell.text.strip() for cell in table.rows[0].cells]
                markdown_content.append("| " + " | ".join(header_cells) + " |")
                markdown_content.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
                
                # Data rows
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    markdown_content.append("| " + " | ".join(cells) + " |")
                
                markdown_content.append("")  # Add blank line after table
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join(markdown_content))
            
            print(f"✓ Successfully converted to Markdown: {docx_file.name}")
            return True
            
        except Exception as e:
            print(f"✗ Error converting {docx_path} to Markdown: {str(e)}")
            return False
    
    def convert_all_pdfs(self) -> dict:
        """
        Convert all PDF files in the input directory
        
        Returns:
            dict: Summary of conversion results
        """
        pdf_files = list(self.input_dir.glob("*.pdf"))
        
        if not pdf_files:
            print(f"No PDF files found in {self.input_dir}")
            return {"total": 0, "successful": 0, "failed": 0}
        
        print(f"\nFound {len(pdf_files)} PDF file(s) to convert\n")
        print("=" * 60)
        
        successful = 0
        failed = 0
        
        for pdf_file in pdf_files:
            if self.convert_single_pdf(str(pdf_file)):
                successful += 1
            else:
                failed += 1
            print("-" * 60)
        
        # Print summary
        print("\n" + "=" * 60)
        print("CONVERSION SUMMARY")
        print("=" * 60)
        print(f"Total files: {len(pdf_files)}")
        print(f"Successful: {successful}")
        print(f"Failed: {failed}")
        print("=" * 60)
        
        return {
            "total": len(pdf_files),
            "successful": successful,
            "failed": failed
        }
    
    def get_pdf_info(self, pdf_path: str) -> dict:
        """
        Get basic information about a PDF file
        
        Args:
            pdf_path: Path to the PDF file
            
        Returns:
            dict: PDF file information
        """
        try:
            pdf_file = Path(pdf_path)
            
            if not pdf_file.exists():
                return {"error": "File not found"}
            
            file_size = pdf_file.stat().st_size
            file_size_mb = file_size / (1024 * 1024)
            
            return {
                "filename": pdf_file.name,
                "path": str(pdf_file.absolute()),
                "size_bytes": file_size,
                "size_mb": round(file_size_mb, 2)
            }
            
        except Exception as e:
            return {"error": str(e)}


def main():
    """Main function to demonstrate usage"""
    print("=" * 60)
    print("PDF to Word Converter")
    print("=" * 60)
    
    # Initialize converter
    converter = PDFToWordConverter()
    
    # Convert all PDFs in input directory
    results = converter.convert_all_pdfs()
    
    if results["total"] == 0:
        print("\nNo PDF files found in the 'input' directory.")
        print("Please add PDF files to the 'input' folder and run again.")


if __name__ == "__main__":
    main()

# Made with Bob
